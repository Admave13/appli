"""
core/doc_editor.py — DOCX: plantillas, copia, lectura y PDF (sin WeasyPrint).
"""

from __future__ import annotations

import os
import re
import shutil
from datetime import datetime

from docx import Document
from docx.shared import Cm, Pt

import config


def get_cv_template_path() -> str:
    path = config.CV_TEMPLATE_PATH
    if not os.path.isfile(path):
        raise FileNotFoundError(f"Plantilla CV no encontrada: {path}")
    return path


def copy_cv_template(dest_path: str) -> str:
    """Copia la plantilla DOCX del CV al destino."""
    os.makedirs(os.path.dirname(dest_path) or ".", exist_ok=True)
    shutil.copy2(get_cv_template_path(), dest_path)
    return dest_path


def copy_cover_template(dest_path: str) -> str:
    os.makedirs(os.path.dirname(dest_path) or ".", exist_ok=True)
    tpl = config.COVER_TEMPLATE_PATH
    if os.path.isfile(tpl):
        shutil.copy2(tpl, dest_path)
    else:
        doc = Document()
        doc.add_paragraph("Estimado equipo de contratación,")
        doc.add_paragraph("")
        doc.add_paragraph("Me dirijo a ustedes para expresar mi interés en el puesto…")
        doc.add_paragraph("")
        doc.add_paragraph("Atentamente,")
        doc.save(dest_path)
    return dest_path


def save_uploaded_docx(uploaded_bytes: bytes, dest_path: str) -> str:
    os.makedirs(os.path.dirname(dest_path) or ".", exist_ok=True)
    with open(dest_path, "wb") as f:
        f.write(uploaded_bytes)
    return dest_path


def text_to_docx(text: str, path: str) -> str:
    """Texto plano → DOCX (fallback; preferir plantilla + subida)."""
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2)
        section.left_margin = section.right_margin = Cm(2.5)
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    doc.save(path)
    return path


def apply_text_to_cv_docx(template_path: str, dest_path: str, new_text: str, keep_header: int = 3) -> str:
    """Copia plantilla y reemplaza cuerpo con texto adaptado (p. ej. IA)."""
    shutil.copy2(template_path, dest_path)
    doc = Document(dest_path)
    body = doc.element.body
    paras = list(doc.paragraphs)
    for p in paras[keep_header:]:
        body.remove(p._element)
    for line in new_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.save(dest_path)
    return dest_path


def docx_to_text(path: str) -> str:
    if not path or not os.path.exists(path):
        return ""
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(lines)


def new_docx_path(directory: str, prefix: str, doc_id: int | None = None) -> str:
    os.makedirs(directory, exist_ok=True)
    suffix = f"_{doc_id}" if doc_id else f"_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in prefix)
    return os.path.join(directory, f"{safe}{suffix}.docx")


def _escape_xml(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def docx_to_pdf(docx_path: str, pdf_path: str) -> str:
    """Convierte DOCX a PDF con ReportLab (funciona en Windows sin GTK)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    doc = Document(docx_path)
    os.makedirs(os.path.dirname(pdf_path) or ".", exist_ok=True)

    pdf = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CvTitle",
        parent=styles["Heading1"],
        fontSize=14,
        leading=16,
        spaceAfter=4,
    )
    heading_style = ParagraphStyle(
        "CvHeading",
        parent=styles["Heading2"],
        fontSize=10,
        leading=12,
        spaceBefore=8,
        spaceAfter=4,
        textColor="#333333",
    )
    body_style = ParagraphStyle(
        "CvBody",
        parent=styles["Normal"],
        fontSize=9,
        leading=11,
        spaceAfter=2,
    )

    story = []
    for i, para in enumerate(doc.paragraphs):
        raw = para.text.strip()
        if not raw:
            story.append(Spacer(1, 4))
            continue
        text = _escape_xml(raw)
        upper_ratio = sum(1 for c in raw if c.isupper()) / max(len(raw), 1)
        is_section = (
            len(raw) < 60
            and upper_ratio > 0.6
            and not re.search(r"\d{3,}", raw)
        )
        if i < 3:
            style = title_style if i == 0 else body_style
        elif is_section:
            style = heading_style
        else:
            style = body_style
        story.append(Paragraph(text, style))
        story.append(Spacer(1, 2))

    pdf.build(story)
    return pdf_path
